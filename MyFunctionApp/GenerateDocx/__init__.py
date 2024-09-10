import os
import uuid
import logging
import azure.functions as func
from docx import Document

def fetch_scan_data_from_cosmos(scan_data):
    return scan_data

def generate_docx_from_knowledge_scan(scan_data):
    doc = Document()

    # Add content to the DOCX file
    doc.add_heading('Knowledge Scan', level=1)
    general_notes = scan_data.get('general_notes', 'No general notes provided.')
    doc.add_heading('General notes:', level=2)
    doc.add_paragraph(general_notes)

    keywords = scan_data.get('keywords', [])
    doc.add_heading('Keywords and themes:', level=3)
    if keywords:
        for keyword in keywords:
            doc.add_paragraph(keyword, style='List Bullet')
    else:
        doc.add_paragraph('No keywords provided.')

    resources_searched = scan_data.get('resources_searched', [])
    doc.add_heading('Resources searched:', level=3)
    if resources_searched:
        for resource in resources_searched:
            doc.add_paragraph(resource, style='List Bullet')
    else:
        doc.add_paragraph('No resources provided.')

    combined_summaries = scan_data.get('combined_summaries', [])
    doc.add_heading('Summaries:', level=2)
    if combined_summaries:
        for i, summary_info in enumerate(combined_summaries, 1):
            doc.add_heading(f"Document {i}: {summary_info['pdf_name']}", level=3)
            doc.add_paragraph(summary_info.get('summary', 'No summary available.'))
    else:
        doc.add_paragraph('No summaries available.')

    # Save the document to a file in the temp directory
    file_name = f"knowledge_scan_{uuid.uuid4()}.docx"
    doc_path = os.path.join('/tmp', file_name)
    doc.save(doc_path)

    return doc_path, file_name

def main(inputDocument: func.DocumentList, outputDocument: func.Out[func.Document], outputBlob: func.Out[str]):
    logging.info("GenerateDocx function triggered.")

    for doc in inputDocument:
        scan_data = fetch_scan_data_from_cosmos(doc)

        # Generate DOCX from the scan data
        doc_path, file_name = generate_docx_from_knowledge_scan(scan_data)

        # Read the content of the DOCX file
        with open(doc_path, 'rb') as file:
            doc_content = file.read()

        # Create a fixed filename for the outputBlob binding
        fixed_blob_path = f"curated/{file_name}"

        # Save the document content to the Blob using a fixed path
        outputBlob.set(doc_content)

        # Prepare the output document item for Cosmos DB
        doc_item = func.Document.from_dict({
            "id": str(uuid.uuid4()),
            "file_name": file_name,
            "blob_location": fixed_blob_path  # Store the Blob storage location in Cosmos DB
        })

        # Output to Cosmos DB docxContainer
        outputDocument.set(doc_item)

    logging.info("DOCX file generated, uploaded to Blob storage, and location saved to Cosmos DB successfully.")
